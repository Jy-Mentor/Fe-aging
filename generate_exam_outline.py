#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
临床实验室管理（第二版）- 简答题与论述题复习提纲
生成Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 78, 121)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    run.bold = True
    return h


def add_bold_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1F4E79')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('临床实验室管理（第二版）')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(28)
    run.bold = True
    p.paragraph_format.space_after = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('简答题与论述题复习提纲')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('满分答案 · 言简意赅版')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(85, 85, 85)
    p.paragraph_format.space_after = Pt(50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('涵盖 12 大核心考点')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ===== 目录页 =====
    add_heading_styled(doc, '目  录', level=1)
    doc.add_paragraph()

    toc_items = [
        ('第一部分  简答题', 1),
        ('  一、室内质量控制（IQC）', 2),
        ('  二、危急值（Critical Value）', 2),
        ('  三、质量管理体系的工作原理', 2),
        ('  四、常见的样本拒收标准', 2),
        ('  五、室内质控物和室间质评物的异同', 2),
        ('  六、PDCA质量改进循环', 2),
        ('  七、临床决定限的概念', 2),
        ('  八、生物安全危害等级的分类', 2),
        ('  九、ROC曲线的应用价值', 2),
        ('  十、室内质控失控处理的一般流程', 2),
        ('  十一、不合格室间质量评价的处理原则', 2),
        ('  十二、Westgard多规则质控原理、步骤和逻辑图', 2),
        ('第二部分  论述题', 1),
        ('  论述题一：质量管理体系的建立与运行', 2),
        ('  论述题二：IQC与EQA的关系及作用', 2),
        ('  论述题三：Westgard多规则质控原理及应用', 2),
        ('  论述题四：PDCA循环在质量改进中的应用', 2),
        ('  论述题五：生物安全管理的重要性及措施', 2),
        ('附录：答题技巧与注意事项', 1),
    ]
    for item, lvl in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12) if lvl == 1 else Pt(11)
        if lvl == 1:
            run.bold = True
        p.paragraph_format.left_indent = Inches(0.3 * (lvl - 1))
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ===== 第一部分 简答题 =====
    add_heading_styled(doc, '第一部分  简答题', level=1)
    doc.add_paragraph()

    # 1. 室内质量控制
    add_heading_styled(doc, '一、室内质量控制（IQC）', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, '室内质量控制是由实验室工作人员采用一定的方法和步骤，连续评价本实验室工作的可靠性程度，确定报告能否发出的一系列活动。')
    add_bold_para(doc, '【目的】')
    add_bullet(doc, '监测实验室测定工作的精密度（日内、日间）')
    add_bullet(doc, '及时发现随机误差和系统误差，采取纠正措施')
    add_bullet(doc, '确保患者检验结果的可靠性和一致性')
    add_bullet(doc, '判断该批次检验报告能否发出')
    add_bold_para(doc, '【核心要素】')
    add_bullet(doc, '质控品：至少2个浓度水平（正常+异常），基质与患者样本相似')
    add_bullet(doc, '质控图：Levey-Jennings质控图（X-S图）最常用')
    add_bullet(doc, '质控规则：Westgard多规则等失控判断标准')
    add_bullet(doc, '失控处理：查找原因→纠正→验证→记录→报告')
    add_bold_para(doc, '【理论基础】')
    add_bullet(doc, '正态分布：±2s覆盖95.5%，±3s覆盖99.7%的测定值')
    add_bullet(doc, '统计学概率：判断误差是否属于偶然误差范围')

    # 2. 危急值
    add_heading_styled(doc, '二、危急值（Critical Value）', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, '危急值是指检验结果极度异常，提示患者可能处于生命危险的边缘状态，临床医师需要及时获得信息并迅速给予干预措施，否则可能危及患者生命的检验数值。')
    add_bold_para(doc, '【报告制度要求】')
    add_bullet(doc, '建立危急值项目表和危急界限值（根据医院等级、科室特点制定）')
    add_bullet(doc, '发现危急值必须立即进行复核（确保结果可靠）')
    add_bullet(doc, '复核无误后立即电话报告临床科室并做好记录')
    add_bullet(doc, '记录内容：患者姓名、性别、年龄、住院号、检验项目、结果、报告时间、报告人、接收人、接收时间')
    add_bold_para(doc, '【常见危急值项目及界限（参考）】')
    add_table(doc,
        ['检验项目', '危急值低限', '危急值高限', '临床意义'],
        [
            ['白细胞计数（WBC）', '＜2.5×10⁹/L', '＞30×10⁹/L', '严重感染/血液系统疾病'],
            ['血红蛋白（Hb）', '＜50g/L', '—', '急性大量失血/严重贫血'],
            ['血小板计数（PLT）', '＜30×10⁹/L', '—', '严重出血倾向'],
            ['血清钾（K⁺）', '＜2.5mmol/L', '＞6.5mmol/L', '严重心律失常/心脏骤停风险'],
            ['血清钠（Na⁺）', '＜120mmol/L', '＞160mmol/L', '严重电解质紊乱'],
            ['血糖（Glu）', '＜2.2mmol/L', '＞22.2mmol/L', '低血糖昏迷/糖尿病酮症酸中毒'],
            ['血液pH', '＜7.20', '＞7.60', '严重酸碱失衡'],
            ['动脉血氧分压', '＜40mmHg', '—', '严重缺氧'],
            ['凝血酶原时间（PT）', '—', '＞30s', '严重出血风险'],
            ['活化部分凝血活酶时间', '—', '＞80s', '严重出血风险'],
        ],
        [4.0, 3.0, 3.0, 4.0]
    )
    add_para(doc, '')
    add_bold_para(doc, '【管理要点】')
    add_bullet(doc, '危急值项目和界限值应由临床和实验室共同制定并定期评审')
    add_bullet(doc, '应有危急值报告登记制度')
    add_bullet(doc, '定期回顾分析危急值报告的及时性和准确性')

    # 3. 质量管理体系的工作原理
    add_heading_styled(doc, '三、质量管理体系的工作原理', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, '质量管理体系是实验室为实现质量方针和质量目标，由组织结构、程序、过程和资源构成的有机整体。')
    add_bold_para(doc, '【八大质量管理原则】')
    add_numbered(doc, '以患者为中心：以满足临床和患者需求为目标')
    add_numbered(doc, '领导作用：领导层确立方向、提供资源')
    add_numbered(doc, '全员参与：各级人员都是组织之本')
    add_numbered(doc, '过程方法：将活动和相关资源作为过程进行管理')
    add_numbered(doc, '管理的系统方法：将相互关联的过程作为系统加以识别、理解和管理')
    add_numbered(doc, '持续改进：是组织永恒的目标（PDCA循环）')
    add_numbered(doc, '基于事实的决策方法：有效决策建立在数据和信息分析的基础上')
    add_numbered(doc, '互利的供方关系：与供方保持互利关系，增强双方创造价值的能力')
    add_bold_para(doc, '【四大组成部分】')
    add_bullet(doc, '组织结构：机构设置、职责权限、相互关系')
    add_bullet(doc, '程序：为进行某项活动所规定的途径（形成文件的程序）')
    add_bullet(doc, '过程：一组将输入转化为输出的相互关联或相互作用的活动')
    add_bullet(doc, '资源：人员、设备、设施、资金、技术方法、信息等')
    add_bold_para(doc, '【核心标准】')
    add_bullet(doc, 'ISO 15189：医学实验室质量和能力的专用要求（最常用）')
    add_bullet(doc, 'ISO/IEC 17025：检测和校准实验室能力的通用要求')
    add_bullet(doc, '医疗机构临床实验室管理办法（卫生部令）')
    add_bold_para(doc, '【过程方法：检验全过程】')
    add_bullet(doc, '分析前：检验申请、患者准备、标本采集、运送、接收')
    add_bullet(doc, '分析中：标本处理、测定、室内质控、结果审核')
    add_bullet(doc, '分析后：结果报告、结果解释、标本保存、咨询服务')
    add_para(doc, '质量管理体系就是对这些全过程进行系统管理，确保每个环节都在受控状态。')

    # 4. 常见的样本拒收标准
    add_heading_styled(doc, '四、常见的样本拒收标准', level=2)
    add_bold_para(doc, '【常见拒收类型及标准】')
    add_table(doc,
        ['拒收类型', '具体情况', '常见受影响项目'],
        [
            ['标本标识错误', '无标签、标签信息不全、与申请单不符', '所有项目（原则性错误）'],
            ['标本量不足', '采血量不足无法完成检测、尿标本量过少', '血常规、凝血功能、生化项目'],
            ['严重溶血', '红细胞破坏释放细胞内成分', '血钾、LDH、AST、胆红素、CK'],
            ['严重脂血', '乳糜血干扰比色和免疫比浊法', '肝功能、血脂、电解质、总蛋白'],
            ['标本凝固', '血液凝固、有凝块（抗凝管）', '血常规、凝血功能、血细胞分析'],
            ['抗凝剂错误', '抗凝管选择错误、抗凝比例不当', '凝血功能（枸橼酸钠1:9）、血常规（EDTA）'],
            ['标本污染', '容器不洁、污染、培养标本污染', '血培养、尿培养、各种微生物培养'],
            ['容器错误', '用错采血管、容器不合格', '对应项目均受影响'],
            ['送检超时', '超过规定时间未送达实验室', '血气、血糖、血氨、凝血功能'],
            ['保存不当', '未避光、未冷藏、未保温等', '胆红素（避光）、血气（冰浴）、冷凝集素（保温）'],
            ['标本泄漏', '容器破损、标本外溢', '受影响项目视情况而定'],
            ['申请单缺失', '无检验申请单或信息无法识别', '所有项目'],
        ],
        [3.0, 5.5, 5.5]
    )
    add_para(doc, '')
    add_bold_para(doc, '【处理原则】')
    add_bullet(doc, '登记拒收原因，通知临床科室重新采集')
    add_bullet(doc, '紧急或特殊情况下（如大出血、抢救等），可先进行检测并在报告中注明异常情况，供临床参考')
    add_bullet(doc, '定期统计分析拒收原因，制定持续改进措施')
    add_bullet(doc, '对医护人员进行标本采集培训，降低不合格率')

    # 5. 室内质控物和室间质评物的异同
    add_heading_styled(doc, '五、室内质控物和室间质评物的异同', level=2)
    add_bold_para(doc, '【相同点】')
    add_bullet(doc, '均为质量保证体系的重要组成部分')
    add_bullet(doc, '均使用已知特性的质控材料')
    add_bullet(doc, '均需按患者标本同样的方式进行测定')
    add_bullet(doc, '均要求材料稳定、均匀、具有互通性')
    add_bullet(doc, '目的都是为了提高检验质量，保证结果可靠')
    add_bold_para(doc, '【不同点】')
    add_table(doc,
        ['比较项目', '室内质控物（IQC）', '室间质评物（EQA/PT）'],
        [
            ['主要目的', '监测精密度、发现随机/系统误差', '评价准确度、发现实验室间系统偏差'],
            ['使用频率', '每批次、每日常规检测', '定期（每月/每季/每年几次）'],
            ['操作者', '本实验室人员', '本实验室人员（盲样检测）'],
            ['结果用途', '实验室内部质量控制', '上报组织者，由组织者评价'],
            ['靶值设定', '本实验室累积均值', '所有实验室均值或参考方法定值'],
            ['评价方式', '质控图+质控规则判断', '与其他实验室比较、打分评价'],
            ['浓度水平', '至少2个水平（正常、异常）', '多个浓度水平，覆盖临床范围'],
            ['基质效应', '可接受（与患者标本相似即可）', '要求高，尽量与临床标本一致'],
            ['结果保密', '实验室内部数据', '组织者保密，返回评价报告'],
            ['发现误差类型', '随机误差、趋势性系统误差', '实验室间系统误差、方法学偏差'],
            ['对临床报告的影响', '直接决定报告能否发出', '回顾性评价，不直接影响单批报告'],
        ],
        [3.0, 5.0, 6.0]
    )
    add_para(doc, '')
    add_bold_para(doc, '【二者关系】')
    add_bullet(doc, '室内质控是室间质评的基础：室内质控不合格，室间质评结果无意义')
    add_bullet(doc, '室间质评是室内质控的补充：可发现室内质控不能发现的系统误差')
    add_bullet(doc, '二者互为补充，缺一不可，共同构成完整的质量保证体系')

    # 6. PDCA质量改进循环
    add_heading_styled(doc, '六、PDCA质量改进循环的工作原理及实施步骤', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, 'PDCA循环又称戴明环（Deming Cycle），是由美国质量管理专家戴明提出的质量管理循环，包括计划（Plan）、执行（Do）、检查（Check）、处理（Action）四个阶段，是持续质量改进的经典方法。')
    add_bold_para(doc, '【四个阶段及实施步骤】')
    add_table(doc,
        ['阶段', '核心任务', '具体工作步骤'],
        [
            ['P（计划）', '分析现状，找出问题，制定改进计划和措施',
             '①分析现状，找出存在的质量问题\n②分析产生问题的各种原因或影响因素\n③找出影响质量的主要因素\n④针对主要因素制定解决措施和行动计划'],
            ['D（执行）', '按计划执行，落实措施',
             '①组织相关人员执行计划\n②按预定的措施和方案落实\n③记录执行过程和数据'],
            ['C（检查）', '对比计划与结果，评估效果',
             '①检查计划执行的结果\n②将实际结果与预期目标对比\n③总结成功经验，找出存在的问题和不足'],
            ['A（处理）', '标准化成功经验，遗留问题进入下一循环',
             '①成功经验加以肯定，纳入标准\n②失败教训加以总结，防止再发\n③未解决的问题转入下一个PDCA循环'],
        ],
        [2.5, 4.0, 7.5]
    )
    add_para(doc, '')
    add_bold_para(doc, '【PDCA循环的特点】')
    add_bullet(doc, '大环套小环：整个组织是一个大PDCA，各部门/科室有各自的小PDCA，层层嵌套')
    add_bullet(doc, '阶梯式上升：每循环一次，质量水平就提高一步，循环不止，改进不断')
    add_bullet(doc, '循环前进：四个阶段周而复始，持续改进永无止境')
    add_bullet(doc, '数据驱动：每个阶段都以事实和数据为依据，用数据说话')
    add_bullet(doc, '全员参与：需要组织内全体成员共同参与')
    add_bold_para(doc, '【临床实验室应用实例】')
    add_bullet(doc, '降低标本不合格率：统计拒收原因→制定培训计划→实施培训→检查效果→标准化')
    add_bullet(doc, '缩短报告周转时间（TAT）：分析延误原因→优化流程→实施→评估效果→改进')
    add_bullet(doc, '提高室间质评成绩：分析不合格原因→制定纠正措施→实施→验证效果→纳入常规')

    # 7. 临床决定限的概念
    add_heading_styled(doc, '七、临床决定限的概念', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, '临床决定限（Clinical Decision Limit）是指检验结果达到某一数值时，临床医生应据此决定对患者采取何种医疗措施的临界值。它是基于临床用途而设定的，不同于参考区间。')
    add_bold_para(doc, '【与参考区间的区别】')
    add_table(doc,
        ['比较项目', '参考区间', '临床决定限'],
        [
            ['定义', '健康人群95%个体的测定值分布范围', '用于临床诊断、治疗、预后决策的阈值'],
            ['设定依据', '健康人群的统计学分布', '疾病诊断标准、治疗指南、预后判断'],
            ['主要目的', '判断结果是否"正常"', '指导临床采取相应医疗措施'],
            ['数量', '通常2个（上限和下限）', '可1个或多个（分级决策）'],
            ['设定方法', '正态分布法、非参数法', 'ROC曲线分析、临床结局研究、专家共识'],
            ['受患病率影响', '不受影响', '部分受疾病流行情况影响'],
            ['与疾病关系', '不直接关联疾病', '直接与疾病诊断/治疗/预后相关'],
        ],
        [3.0, 5.0, 6.0]
    )
    add_para(doc, '')
    add_bold_para(doc, '【临床决定限的类型】')
    add_numbered(doc, '诊断决定限：用于疾病的诊断和排除')
    add_bullet(doc, '例：糖尿病诊断：空腹血糖≥7.0mmol/L', 1)
    add_bullet(doc, '例：急性心肌梗死：肌钙蛋白＞第99百分位值', 1)
    add_numbered(doc, '治疗决定限：决定是否开始治疗或调整治疗方案')
    add_bullet(doc, '例：血脂异常：LDL-C目标值根据心血管风险分层', 1)
    add_bullet(doc, '例：高血压：血压≥140/90mmHg考虑药物治疗', 1)
    add_numbered(doc, '预后决定限：判断疾病预后和风险分层')
    add_bullet(doc, '例：肿瘤标志物水平与预后的关系', 1)
    add_numbered(doc, '危急值：见危急值章节，是特殊的临床决定限')
    add_bold_para(doc, '【设定方法】')
    add_bullet(doc, 'ROC曲线分析：选择最佳截断值（cut-off值），平衡敏感度和特异度')
    add_bullet(doc, '临床结局研究：前瞻性观察不同水平患者的临床结局')
    add_bullet(doc, '专家共识和指南推荐：由权威机构制定')
    add_bullet(doc, '根据临床需求调整：根据筛查、诊断、监测等不同目的选择不同切点')

    # 8. 生物安全危害等级的分类
    add_heading_styled(doc, '八、生物安全危害等级的分类', level=2)
    add_bold_para(doc, '【分类依据】')
    add_para(doc, '根据微生物的致病性、感染剂量、传播途径、宿主范围、有效治疗和预防措施的可用性等因素，将生物因子分为四个危害等级。')
    add_bold_para(doc, '【四级分类标准】')
    add_table(doc,
        ['等级', '名称', '危害特点', '防护级别', '典型病原体'],
        [
            ['Ⅰ级', '低个体危害，低群体危害', '不会导致健康工作者和动物致病的细菌、真菌、病毒和寄生虫等生物因子',
             'BSL-1\n基础实验室', '大肠杆菌（非致病性）、枯草杆菌、啤酒酵母'],
            ['Ⅱ级', '中等个体危害，有限群体危害', '能引起人或动物发病，但一般情况下对健康工作者、群体、家畜或环境不会引起严重危害的病原体；有有效治疗和预防措施',
             'BSL-2\n基础实验室', '金黄色葡萄球菌、乙肝病毒、流感病毒、沙门菌、志贺菌、梅毒螺旋体'],
            ['Ⅲ级', '高个体危害，低群体危害', '能引起人或动物严重疾病，或造成严重经济损失，但通常不能因偶然接触而在个体间传播，或能使用抗生素、抗寄生虫药治疗的病原体',
             'BSL-3\n防护实验室', '结核分枝杆菌、HIV、SARS-CoV、布鲁氏菌、炭疽杆菌、狂犬病毒'],
            ['Ⅳ级', '高个体危害，高群体危害', '能引起人或动物非常严重的疾病，一般不能治愈，容易直接间接或因偶然接触在人与人、动物与人、人与动物、动物与动物间传播的病原体',
             'BSL-4\n最高防护实验室', '埃博拉病毒、马尔堡病毒、天花病毒、克里米亚-刚果出血热病毒'],
        ],
        [1.0, 2.8, 4.0, 2.0, 4.2]
    )
    add_para(doc, '')
    add_bold_para(doc, '【实验室生物安全防护分级及要求】')
    add_numbered(doc, 'BSL-1（一级生物安全防护实验室）')
    add_bullet(doc, '操作要求：标准微生物操作，开放式工作台', 1)
    add_bullet(doc, '设施要求：普通实验室，有洗手池', 1)
    add_numbered(doc, 'BSL-2（二级生物安全防护实验室）')
    add_bullet(doc, '操作要求：生物安全柜处理感染性材料，个人防护装备', 1)
    add_bullet(doc, '设施要求：生物安全柜、高压灭菌器、洗眼装置', 1)
    add_numbered(doc, 'BSL-3（三级生物安全防护实验室）')
    add_bullet(doc, '操作要求：所有感染性材料必须在生物安全柜内操作', 1)
    add_bullet(doc, '设施要求：负压、双门进入、空气过滤排出、独立通风系统', 1)
    add_numbered(doc, 'BSL-4（四级生物安全防护实验室）')
    add_bullet(doc, '操作要求：正压防护服、完全隔离、双人操作', 1)
    add_bullet(doc, '设施要求：完全隔离建筑、独立供气和排气系统、双高压灭菌器', 1)

    # 9. ROC曲线的应用价值
    add_heading_styled(doc, '九、ROC曲线的应用价值', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, 'ROC曲线（受试者工作特征曲线，Receiver Operating Characteristic Curve）是以假阳性率（1-特异度）为横坐标，真阳性率（敏感度）为纵坐标绘制的曲线，用于全面、准确评价诊断试验的准确性。')
    add_bold_para(doc, '【主要应用价值】')
    add_numbered(doc, '评价诊断试验的准确性')
    add_bullet(doc, 'AUC（曲线下面积）越大，诊断准确性越高', 1)
    add_bullet(doc, 'AUC = 0.5：无诊断价值（与随机猜测相同）', 1)
    add_bullet(doc, 'AUC 0.5~0.7：较低准确性', 1)
    add_bullet(doc, 'AUC 0.7~0.9：中等准确性', 1)
    add_bullet(doc, 'AUC 0.9~1.0：较高准确性', 1)
    add_numbered(doc, '确定最佳截断值（cut-off值）')
    add_bullet(doc, '根据临床需求选择适当的切点', 1)
    add_bullet(doc, '约登指数（Youden Index）最大法：敏感度+特异度-1 最大', 1)
    add_bullet(doc, '可根据漏诊/误诊的代价调整', 1)
    add_numbered(doc, '比较两种或多种诊断试验的优劣')
    add_bullet(doc, '比较AUC的统计学差异', 1)
    add_bullet(doc, '同一组受试者中比较，消除人群偏倚', 1)
    add_numbered(doc, '设定临床决定限')
    add_bullet(doc, '为临床决策提供循证依据', 1)
    add_bullet(doc, '可绘制不同患病率下的预测值曲线', 1)
    add_bold_para(doc, '【ROC曲线的优点】')
    add_bullet(doc, '不受患病率影响，客观反映诊断试验本身的性能')
    add_bullet(doc, '全面展示敏感度与特异度的关系（所有截断值的情况）')
    add_bullet(doc, '可用于定量资料和等级资料的诊断准确性评价')
    add_bullet(doc, '直观、形象，易于理解和交流')
    add_bullet(doc, '通过AUC可进行不同诊断试验的量化比较')
    add_bold_para(doc, '【绘制方法】')
    add_bullet(doc, '金标准确认患病组和非患病组')
    add_bullet(doc, '用待评价方法检测两组人群')
    add_bullet(doc, '设定不同截断值，计算对应的敏感度和1-特异度')
    add_bullet(doc, '以1-特异度为横坐标，敏感度为纵坐标作图')

    # 10. 室内质控失控处理的一般流程
    add_heading_styled(doc, '十、室内质控失控处理的一般流程', level=2)
    add_bold_para(doc, '【总原则】')
    add_para(doc, '立即暂停报告发出，查找失控原因，采取纠正措施，验证有效后方可重新报告。')
    add_bold_para(doc, '【失控处理步骤】')
    add_numbered(doc, '立即暂停该批次检验结果的发出，启动失控处理程序')
    add_numbered(doc, '第一步：立即重测同一质控品')
    add_bullet(doc, '目的：排除偶然误差和操作失误', 1)
    add_bullet(doc, '重测结果在控→可能是随机误差，可发出报告，继续观察', 1)
    add_bullet(doc, '重测仍失控→进入下一步', 1)
    add_numbered(doc, '第二步：新开一瓶质控品重测')
    add_bullet(doc, '目的：排除质控品变质、污染、溶解不当的问题', 1)
    add_bullet(doc, '新瓶在控→原质控品问题，更换质控品', 1)
    add_bullet(doc, '新瓶仍失控→进入下一步', 1)
    add_numbered(doc, '第三步：检查仪器和校准状态')
    add_bullet(doc, '检查仪器状态：温度、压力、光源、比色杯等', 1)
    add_bullet(doc, '执行仪器维护：清洗、保养', 1)
    add_bullet(doc, '重新校准仪器：检查校准品、校准曲线', 1)
    add_numbered(doc, '第四步：更换试剂重测')
    add_bullet(doc, '目的：排除试剂变质、配制错误、批号更换等问题', 1)
    add_bullet(doc, '新试剂在控→原试剂问题，更换试剂', 1)
    add_bullet(doc, '新试剂仍失控→进入下一步', 1)
    add_numbered(doc, '第五步：寻求技术支持')
    add_bullet(doc, '联系厂家工程师或试剂供应商技术支持', 1)
    add_bullet(doc, '请教实验室资深技术人员', 1)
    add_numbered(doc, '第六步：问题解决后的验证')
    add_bullet(doc, '重新测定质控品，确认在控（连续2-3次）', 1)
    add_bullet(doc, '对失控期间的患者标本进行评估和必要的复测', 1)
    add_numbered(doc, '第七步：记录与总结')
    add_bullet(doc, '填写失控报告，记录失控规则、原因、纠正措施、验证结果', 1)
    add_bullet(doc, '定期分析失控原因，制定预防措施（PDCA持续改进）', 1)
    add_bold_para(doc, '【失控原因分类】')
    add_bullet(doc, '随机误差：偶然因素，单次失控，重测后恢复，如操作失误、电源波动等')
    add_bullet(doc, '系统误差：连续同向偏移（均值漂移），由仪器、试剂、校准品等因素引起')
    add_bullet(doc, '精密度变化：标准差逐渐扩大，由仪器性能下降引起')

    # 11. 不合格室间质量评价的处理原则
    add_heading_styled(doc, '十一、不合格室间质量评价的处理原则', level=2)
    add_bold_para(doc, '【定义】')
    add_para(doc, '室间质量评价（EQA/PT）成绩不合格是指实验室某次室间质评结果未达到组织者规定的合格标准，提示实验室可能存在系统误差或其他质量问题。')
    add_bold_para(doc, '【处理的总原则】')
    add_bullet(doc, '立即启动纠正和预防措施（CAPA）')
    add_bullet(doc, '查找根本原因，不流于表面')
    add_bullet(doc, '采取纠正措施并验证效果')
    add_bullet(doc, '形成完整的调查处理记录和报告')
    add_bullet(doc, '持续改进，防止再次发生')
    add_bold_para(doc, '【处理步骤】')
    add_numbered(doc, '收到不合格通知后，立即组织分析调查')
    add_numbered(doc, '检查原始记录：标本接收、保存、处理、测定、审核全过程')
    add_numbered(doc, '检查同期室内质控情况：')
    add_bullet(doc, '同期室内质控也失控→提示存在真实质量问题', 1)
    add_bullet(doc, '同期室内质控在控→可能是操作误差、基质效应等', 1)
    add_numbered(doc, '分析可能原因（按可能性大小排查）：')
    add_bullet(doc, '方法学问题：方法学本身缺陷、试剂质量问题、校准品问题', 1)
    add_bullet(doc, '仪器问题：校准偏差、性能下降、维护不当', 1)
    add_bullet(doc, '操作问题：人员操作不规范、计算错误、抄写错误', 1)
    add_bullet(doc, '标本问题：处理不当、保存不当、复溶错误', 1)
    add_bullet(doc, '室间质评物问题：基质效应、不均匀、稳定性差（需证据支持）', 1)
    add_numbered(doc, '制定并实施纠正措施')
    add_numbered(doc, '验证纠正效果：')
    add_bullet(doc, '重新检测留存的室间质评标本', 1)
    add_bullet(doc, '参加后续的再评价或补充评价', 1)
    add_bullet(doc, '与其他实验室比对（注意合规）', 1)
    add_bullet(doc, '用有证标准物质验证', 1)
    add_numbered(doc, '形成完整的不合格调查和处理报告')
    add_numbered(doc, '纳入预防措施，防止类似问题再次发生')
    add_bold_para(doc, '【注意事项】')
    add_bullet(doc, '禁止与其他实验室核对室间质评结果（违规行为）')
    add_bullet(doc, '必须由本实验室独立完成调查和纠正过程')
    add_bullet(doc, '连续多次不合格可能导致实验室认可暂停或撤销')
    add_bullet(doc, '调查处理记录应妥善保存，以备检查')

    # 12. Westgard多规则质控
    add_heading_styled(doc, '十二、Westgard多规则质控原理、步骤和逻辑图', level=2)
    add_bold_para(doc, '【基本原理】')
    add_para(doc, 'Westgard多规则质控是将多个统计学质控规则联合应用，利用不同规则对不同类型误差敏感性不同的特点，既提高误差检出率（真失控率），又控制假失控率在较低水平。')
    add_bold_para(doc, '【常用6条核心规则（2个浓度水平）】')
    add_table(doc,
        ['规则符号', '规则含义', '误差类型', '临床意义'],
        [
            ['1₂s', '1个质控值超出±2s（警告规则）', '随机/系统', '仅作为警告，启动检查程序，不判定失控'],
            ['1₃s', '1个质控值超出±3s', '随机误差', '提示存在较大随机误差，或严重的系统误差'],
            ['2₂s', '连续2个质控值在同一侧超出±2s', '系统误差', '提示均值漂移，存在系统误差'],
            ['R₄s', '同批次2个浓度质控值之差＞4s', '随机误差', '提示精密度变差，随机误差增大'],
            ['4₁s', '连续4个质控值在同一侧超出±1s', '系统误差', '提示较小的均值漂移，早期发现系统误差'],
            ['10x̄', '连续10个质控值落在均值同一侧', '系统误差', '提示持续的小偏差，系统趋势变化'],
        ],
        [2.0, 5.0, 2.5, 4.5]
    )
    add_para(doc, '')
    add_bold_para(doc, '【判断步骤（逻辑流程）】')
    add_numbered(doc, '测定2个浓度水平质控品（高值H、低值L）')
    add_numbered(doc, '检查1₂s规则：任一质控值是否超出±2s')
    add_bullet(doc, '否→在控，可以发出报告', 1)
    add_bullet(doc, '是→警告状态，继续检查以下规则', 1)
    add_numbered(doc, '检查1₃s规则：有无超出±3s的质控值')
    add_bullet(doc, '有→失控（随机误差）', 1)
    add_bullet(doc, '无→继续检查', 1)
    add_numbered(doc, '检查2₂s规则：')
    add_bullet(doc, '同浓度连续2次同侧超2s（如H连续2次+2s）', 1)
    add_bullet(doc, '同批2个浓度同侧超2s（如H和L均+2s）', 1)
    add_bullet(doc, '有→失控（系统误差）', 1)
    add_bullet(doc, '无→继续检查', 1)
    add_numbered(doc, '检查R₄s规则：同批高值与低值之差是否＞4s')
    add_bullet(doc, '是→失控（随机误差）', 1)
    add_bullet(doc, '否→继续检查', 1)
    add_numbered(doc, '检查4₁s规则：连续4个质控值是否同侧超1s')
    add_bullet(doc, '是→失控（系统误差）', 1)
    add_bullet(doc, '否→继续检查', 1)
    add_numbered(doc, '检查10x̄规则：连续10个质控值是否在均值同侧')
    add_bullet(doc, '是→失控（系统误差）', 1)
    add_bullet(doc, '否→在控（1₂s仅为警告），可发出报告', 1)
    add_bold_para(doc, '【逻辑图（文字版）】')
    add_para(doc, '')
    add_para(doc, '            ┌──────────────┐')
    add_para(doc, '            │  测定质控品  │')
    add_para(doc, '            └──────┬───────┘')
    add_para(doc, '                   ↓')
    add_para(doc, '            ┌──────────────┐')
    add_para(doc, '            │  检查 1₂s    │──无超限──→ 在控，发报告')
    add_para(doc, '            └──────┬───────┘')
    add_para(doc, '                 有超限')
    add_para(doc, '                   ↓')
    add_para(doc, '    ┌─────────────────────────────┐')
    add_para(doc, '    │ 1₃s ?  2₂s ?  R₄s ?         │──任一满足──→ 失控')
    add_para(doc, '    └──────────┬──────────────────┘')
    add_para(doc, '               均不满足')
    add_para(doc, '                   ↓')
    add_para(doc, '    ┌─────────────────────────────┐')
    add_para(doc, '    │   4₁s ?   10x̄ ?             │──任一满足──→ 失控')
    add_para(doc, '    └──────────┬──────────────────┘')
    add_para(doc, '               均不满足')
    add_para(doc, '                   ↓')
    add_para(doc, '          在控（1₂s警告）')
    add_para(doc, '')
    add_bold_para(doc, '【Westgard多规则的优点】')
    add_bullet(doc, '低假失控率：约0.3%~3%，比单用规则更合理')
    add_bullet(doc, '高误差检出率：对随机误差和系统误差均有较高的检出能力')
    add_bullet(doc, '误差分型：可初步判断误差类型（随机vs系统），指导失控原因分析')
    add_bullet(doc, '适用性广：适合各种临床检验项目')
    add_bullet(doc, '国际通用：结果可比性好，便于交流和认可')
    add_bold_para(doc, '【注意事项】')
    add_bullet(doc, '需要至少2个浓度水平的质控品')
    add_bullet(doc, '需要累积足够的数据建立稳定的均值和标准差')
    add_bullet(doc, '操作人员需要经过培训，正确理解和应用规则')
    add_bullet(doc, '应根据项目特点选择合适的规则组合')

    doc.add_page_break()

    # ===== 第二部分 论述题 =====
    add_heading_styled(doc, '第二部分  论述题', level=1)
    doc.add_paragraph()

    # 论述题一
    add_heading_styled(doc, '论述题一：试述临床实验室质量管理体系的建立与运行', level=2)
    add_bold_para(doc, '【答题框架】')
    add_numbered(doc, '质量管理体系的概念与意义')
    add_bullet(doc, '定义：为实现质量方针和目标，由组织结构、程序、过程、资源构成的有机整体', 1)
    add_bullet(doc, '意义：保证检验质量、满足临床需求、提高管理水平、是实验室认可的基础', 1)
    add_numbered(doc, '建立质量管理体系的步骤')
    add_bullet(doc, '领导层决策与准备：统一认识、确定目标、提供资源', 1)
    add_bullet(doc, '现状分析与差距评估：对照标准（ISO 15189）进行差距分析', 1)
    add_bullet(doc, '制定质量方针和质量目标：结合实验室实际，切实可行', 1)
    add_bullet(doc, '组织结构设计与职责分配：明确各岗位职责权限', 1)
    add_bullet(doc, '体系文件编写：质量手册、程序文件、作业指导书、记录表格四级文件', 1)
    add_bullet(doc, '体系发布与试运行：全员培训、试运行、发现问题并改进', 1)
    add_numbered(doc, '质量管理体系的运行')
    add_bullet(doc, '全员培训：各级人员理解并掌握体系要求', 1)
    add_bullet(doc, '严格按文件执行：所有工作有据可依、有章可循', 1)
    add_bullet(doc, '做好质量记录：为体系运行提供客观证据', 1)
    add_bullet(doc, '内部审核：定期检查体系运行的符合性和有效性', 1)
    add_bullet(doc, '管理评审：领导层对体系的适宜性、充分性、有效性进行评审', 1)
    add_bullet(doc, '持续改进：通过PDCA循环不断完善体系', 1)
    add_numbered(doc, '质量控制与质量保证的关系')
    add_bullet(doc, '室内质控：日常精密度监测', 1)
    add_bullet(doc, '室间质评：准确度评价', 1)
    add_bullet(doc, '仪器校准：量值溯源', 1)
    add_bullet(doc, '人员培训：能力保证', 1)
    add_numbered(doc, '临床实验室认可的作用（ISO 15189）')
    add_bullet(doc, '提高实验室管理水平和技术能力', 1)
    add_bullet(doc, '增强临床和患者的信任', 1)
    add_bullet(doc, '促进国际交流与结果互认', 1)

    # 论述题二
    add_heading_styled(doc, '论述题二：试述室内质量控制与室间质量评价的关系及在质量保证中的作用', level=2)
    add_bold_para(doc, '【答题框架】')
    add_numbered(doc, '室内质量控制（IQC）概述')
    add_bullet(doc, '概念：由实验室内部人员进行的质量控制活动', 1)
    add_bullet(doc, '目的：监测精密度，发现误差，确保报告可靠', 1)
    add_bullet(doc, '方法：质控图、质控规则（如Westgard多规则）', 1)
    add_bullet(doc, '特点：每日常规、连续监测、实时判断', 1)
    add_numbered(doc, '室间质量评价（EQA/PT）概述')
    add_bullet(doc, '概念：由外部组织者组织的实验室间比对活动', 1)
    add_bullet(doc, '目的：评价准确度，发现系统误差，提高可比性', 1)
    add_bullet(doc, '方法：发放盲样、实验室检测、组织者评价', 1)
    add_bullet(doc, '特点：定期开展、回顾性评价、横向比较', 1)
    add_numbered(doc, '二者的比较（异同点）')
    add_table(doc,
        ['比较项目', '室内质控', '室间质评'],
        [
            ['评价内容', '精密度（重复性）', '准确度（与其他实验室比较）'],
            ['误差类型', '随机误差、趋势性系统误差', '实验室间系统误差'],
            ['频率', '每批次/每日', '定期（每月/每季/每年）'],
            ['作用时机', '实时，直接影响报告发出', '回顾性，事后评价'],
            ['结果对比', '与本实验室历史数据比', '与其他实验室或靶值比'],
        ],
        [3.0, 5.0, 6.0]
    )
    add_para(doc, '')
    add_numbered(doc, '二者的关系')
    add_bullet(doc, '室内质控是室间质评的基础：室内质控失控的结果不能用于室间质评；室内质控良好是室间质评合格的前提', 1)
    add_bullet(doc, '室间质评是室内质控的补充：可发现室内质控不能发现的系统误差（方法学偏差、校准偏差等）', 1)
    add_bullet(doc, '二者互为补充，缺一不可，共同构成完整的质量保证体系', 1)
    add_numbered(doc, '在质量保证体系中的作用')
    add_bullet(doc, '质量保证（QA）是一个更大的概念，涵盖IQC和EQA', 1)
    add_bullet(doc, 'IQC确保日内、日间精密度，控制常规工作质量', 1)
    add_bullet(doc, 'EQA确保实验室间结果的可比性和准确性', 1)
    add_bullet(doc, 'IQC是基础，EQA是验证和补充', 1)
    add_bullet(doc, '共同保障检验结果的准确、可靠、可比', 1)

    # 论述题三
    add_heading_styled(doc, '论述题三：试述Westgard多规则质控的原理、组成及临床应用价值', level=2)
    add_bold_para(doc, '【答题框架】')
    add_numbered(doc, 'Westgard多规则质控的基本原理')
    add_bullet(doc, '基于统计学原理，利用正态分布规律判断误差', 1)
    add_bullet(doc, '联合应用多个质控规则，发挥各规则的优势', 1)
    add_bullet(doc, '不同规则对不同类型误差敏感性不同：1₃s、R₄s对随机误差敏感；2₂s、4₁s、10x̄对系统误差敏感', 1)
    add_bullet(doc, '1₂s作为警告规则启动检查，其他规则判断失控', 1)
    add_bullet(doc, '目标：高误差检出率 + 低假失控率', 1)
    add_numbered(doc, '常用规则及意义（6条核心规则）')
    add_bullet(doc, '1₂s：警告规则，超出±2s即进入检查程序', 1)
    add_bullet(doc, '1₃s：1个值超出±3s，提示随机误差', 1)
    add_bullet(doc, '2₂s：连续2个值同侧超2s，提示系统误差', 1)
    add_bullet(doc, 'R₄s：同批高低值之差＞4s，提示随机误差', 1)
    add_bullet(doc, '4₁s：连续4个值同侧超1s，提示系统误差', 1)
    add_bullet(doc, '10x̄：连续10个值在均值同侧，提示系统误差', 1)
    add_numbered(doc, '判断逻辑与步骤')
    add_bullet(doc, '先查1₂s→无超限→在控', 1)
    add_bullet(doc, '有超限→依次查1₃s、2₂s、R₄s、4₁s、10x̄', 1)
    add_bullet(doc, '任一规则满足→失控', 1)
    add_bullet(doc, '均不满足→在控（警告）', 1)
    add_numbered(doc, '临床应用价值')
    add_bullet(doc, '提高质量控制的准确性和可靠性：比单规则更科学', 1)
    add_bullet(doc, '误差分型：可初步判断随机或系统误差，指导失控原因分析', 1)
    add_bullet(doc, '假失控率低：约0.3%~3%，减少不必要的重复工作', 1)
    add_bullet(doc, '国际通用：结果可比性好，便于交流和认可', 1)
    add_bullet(doc, '可根据项目特点调整规则组合', 1)
    add_numbered(doc, '局限性与注意事项')
    add_bullet(doc, '需要至少2个浓度水平质控品', 1)
    add_bullet(doc, '需要累积足够数据建立稳定的均值和标准差', 1)
    add_bullet(doc, '人员需要培训，正确理解和应用规则', 1)
    add_bullet(doc, '不能替代全面的质量管理体系', 1)

    # 论述题四
    add_heading_styled(doc, '论述题四：试述PDCA循环在临床实验室持续质量改进中的应用', level=2)
    add_bold_para(doc, '【答题框架】')
    add_numbered(doc, 'PDCA循环的概念和起源')
    add_bullet(doc, '又称戴明环，由美国质量管理专家戴明提出', 1)
    add_bullet(doc, '包括Plan（计划）、Do（执行）、Check（检查）、Action（处理）四个阶段', 1)
    add_bullet(doc, '是持续质量改进的经典方法', 1)
    add_numbered(doc, 'PDCA四个阶段的具体内容')
    add_bullet(doc, 'Plan（计划）：分析现状找问题→分析原因→找主因→定计划', 1)
    add_bullet(doc, 'Do（执行）：按计划执行，落实措施，记录过程', 1)
    add_bullet(doc, 'Check（检查）：检查效果，与目标对比，总结经验教训', 1)
    add_bullet(doc, 'Action（处理）：成功经验标准化，失败教训防止再发，遗留问题进入下一循环', 1)
    add_numbered(doc, 'PDCA循环的特点')
    add_bullet(doc, '大环套小环：组织大PDCA，部门小PDCA，层层嵌套', 1)
    add_bullet(doc, '阶梯式上升：每循环一次质量提高一步', 1)
    add_bullet(doc, '循环前进：持续改进，永无止境', 1)
    add_bullet(doc, '数据驱动：以事实和数据为依据', 1)
    add_numbered(doc, '在临床实验室中的应用实例')
    add_bullet(doc, '降低样本拒收率：统计拒收原因→制定培训计划→实施培训→检查效果→标准化', 1)
    add_bullet(doc, '缩短报告周转时间（TAT）：分析延误原因→优化流程→实施→评估→改进', 1)
    add_bullet(doc, '提高室间质评成绩：分析不合格原因→制定纠正措施→实施→验证→纳入常规', 1)
    add_bullet(doc, '降低室内质控失控率：分析失控原因→制定预防措施→实施→验证→改进', 1)
    add_bullet(doc, '提高患者满意度：调查满意度→找问题→改措施→再评价→持续改进', 1)
    add_numbered(doc, '实施的关键要素')
    add_bullet(doc, '领导重视：管理层推动和支持', 1)
    add_bullet(doc, '全员参与：各级人员积极参与', 1)
    add_bullet(doc, '数据驱动：用数据说话，实事求是', 1)
    add_bullet(doc, '制度保障：建立长效机制', 1)
    add_bullet(doc, '持续投入：资源、培训、时间', 1)

    # 论述题五
    add_heading_styled(doc, '论述题五：试述生物安全管理在临床实验室中的重要性及措施', level=2)
    add_bold_para(doc, '【答题框架】')
    add_numbered(doc, '生物安全的概念和重要性')
    add_bullet(doc, '概念：避免危险生物因子造成实验室人员暴露、向实验室外扩散并导致危害的综合措施', 1)
    add_bullet(doc, '保护实验室工作人员的健康和安全', 1)
    add_bullet(doc, '防止医院感染和社区感染', 1)
    add_bullet(doc, '保护环境，防止病原微生物扩散', 1)
    add_bullet(doc, '法律法规要求（《生物安全法》《病原微生物实验室生物安全管理条例》）', 1)
    add_numbered(doc, '生物危害分级')
    add_bullet(doc, 'Ⅰ级：低个体危害，低群体危害', 1)
    add_bullet(doc, 'Ⅱ级：中等个体危害，有限群体危害（临床实验室常见）', 1)
    add_bullet(doc, 'Ⅲ级：高个体危害，低群体危害', 1)
    add_bullet(doc, 'Ⅳ级：高个体危害，高群体危害', 1)
    add_numbered(doc, '实验室生物安全防护分级')
    add_bullet(doc, 'BSL-1：基础实验室，一级屏障', 1)
    add_bullet(doc, 'BSL-2：基础实验室，生物安全柜+个人防护（临床微生物、HIV初筛等）', 1)
    add_bullet(doc, 'BSL-3：防护实验室，负压、双门、空气过滤', 1)
    add_bullet(doc, 'BSL-4：最高防护实验室，完全隔离', 1)
    add_numbered(doc, '生物安全管理措施')
    add_bullet(doc, '管理体系：成立生物安全委员会，制定管理制度和操作规程', 1)
    add_bullet(doc, '人员管理：培训、考核、持证上岗、健康监护', 1)
    add_bullet(doc, '设施设备：生物安全柜、高压灭菌器、洗眼器、应急喷淋', 1)
    add_bullet(doc, '标本管理：接收登记、转运防护、处理规范、保存安全', 1)
    add_bullet(doc, '菌（毒）种管理：专人负责、双人双锁、使用登记、销毁记录', 1)
    add_bullet(doc, '废弃物处理：分类收集、消毒灭菌、标识清楚、记录完整', 1)
    add_bullet(doc, '意外事故处理：应急预案、暴露后处置流程、定期演练', 1)
    add_numbered(doc, '临床实验室常见生物安全风险及防控')
    add_bullet(doc, '气溶胶吸入：生物安全柜内操作、规范离心', 1)
    add_bullet(doc, '针刺伤：使用安全器具、规范操作、禁止回套针帽', 1)
    add_bullet(doc, '样本溅洒：立即消毒处理程序', 1)
    add_bullet(doc, '职业暴露后处理：紧急处理→报告→评估→预防用药→随访', 1)

    doc.add_page_break()

    # ===== 附录 =====
    add_heading_styled(doc, '附录：答题技巧与注意事项', level=1)
    doc.add_paragraph()

    add_heading_styled(doc, '一、简答题答题技巧', level=2)
    add_numbered(doc, '定义先行：先答定义/概念，再展开要点')
    add_numbered(doc, '条理清晰：用序号分点作答，层次分明')
    add_numbered(doc, '言简意赅：答要点即可，无需过度展开')
    add_numbered(doc, '关键词准确：专业术语使用正确')
    add_numbered(doc, '看分答题：分值高的多答，分值低的少答')

    add_heading_styled(doc, '二、论述题答题技巧', level=2)
    add_numbered(doc, '结构完整：采用总-分-总结构，有开头有结尾')
    add_numbered(doc, '层次分明：分段论述，每段一个核心观点')
    add_numbered(doc, '内容全面：覆盖题目所有方面，不遗漏要点')
    add_numbered(doc, '逻辑严谨：前后呼应，论证充分')
    add_numbered(doc, '结合实际：适当联系临床实验室实际工作')
    add_numbered(doc, '字数适宜：根据分值决定篇幅，一般不少于400字')

    add_heading_styled(doc, '三、考试注意事项', level=2)
    add_numbered(doc, '仔细审题：明确题目要求，避免答非所问')
    add_numbered(doc, '合理分配时间：根据题型和分值安排答题时间')
    add_numbered(doc, '字迹工整：保持卷面整洁')
    add_numbered(doc, '先易后难：先答会的，再回头考虑难题')
    add_numbered(doc, '专业术语准确：体现专业素养')

    # 保存文档
    output_path = r'D:\铁衰老 绝不重蹈覆辙\临床实验室管理_简答题论述题提纲.docx'
    doc.save(output_path)
    print(f'文档生成成功：{output_path}')


if __name__ == '__main__':
    main()
