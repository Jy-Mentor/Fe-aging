#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def h1(doc, text):
    h = doc.add_heading(level=1)
    r = h.add_run(text)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.bold = True
    return h


def h2(doc, text):
    h = doc.add_heading(level=2)
    r = h.add_run(text)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r.font.size = Pt(14)
    r.bold = True
    return h


def bp(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)
    return p


def blt(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    return p


def num(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    return p


def mktable(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(c, '1F4E79')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(10)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t
